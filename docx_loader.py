from docx import Document as DocxDocument
from langchain_core.documents import Document
import os
import logging
from utils import setup_logger

logger = setup_logger("DocxLoader")

def load_docx(file_path: str) -> list[Document]:
    """
    Extract text from DOCX files.
    """
    file_name = os.path.basename(file_path)
    documents = []
    
    try:
        doc = DocxDocument(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        metadata = {
            "source": file_name,
            "file_name": file_name,
            "page_number": 1,
            "file_type": "docx",
            "content_type": "text",
            "method": "python-docx"
        }
        
        content = "\n".join(full_text)
        if content:
            documents.append(Document(page_content=content, metadata=metadata))
            logger.info(f"Loaded DOCX: {file_name}")
        
    except Exception as e:
        logger.error(f"Error loading DOCX {file_path}: {e}")
        
    return documents
